from selenium import webdriver
from selenium.webdriver.common.by import By
import time
from docx import Document
from docx.shared import Inches
import openpyxl
from PIL import Image



driver = webdriver.Chrome()
driver.maximize_window()

driver.get("https://exam.smvec.ac.in/exam_result_ug_pg_aug_sep_2023/")
time.sleep(3)
doc = Document()


def read_captcha():
    captcha =  driver.find_element("xpath","/html/body/div[3]/div[1]/div/div/div[2]/div[2]/div[2]/div/div/div/div/form/div[3]/div/span[1]")
    return captcha.text

def get_result(regno_value,dob_value):
    regno = driver.find_element("xpath","/html/body/div[3]/div[1]/div/div/div[2]/div[2]/div[2]/div/div/div/div/form/div[1]/div/input")
    regno.send_keys(regno_value)

    dob = driver.find_element("xpath","/html/body/div[3]/div[1]/div/div/div[2]/div[2]/div[2]/div/div/div/div/form/div[2]/div/input")
    dob.send_keys(dob_value)

    captcha = driver.find_element("xpath","/html/body/div[3]/div[1]/div/div/div[2]/div[2]/div[2]/div/div/div/div/form/div[3]/div/input")
    captcha.send_keys(read_captcha())
    
    submit = driver.find_element("xpath","/html/body/div[3]/div[1]/div/div/div[2]/div[2]/div[2]/div/div/div/div/form/div[4]/div/button[1]")
    submit.click()

def read_result():
    result = dict()
    table = driver.find_element(By.XPATH, "/html/body/div[3]/div[1]/div/div/div[3]/div[2]/div/div/div/div/div/div/div[6]/table")
    rows = table.find_elements(By.TAG_NAME, "tr")
    for row in rows[1:]:  
        columns = row.find_elements(By.TAG_NAME, "td")
        if len(columns) == 2:  
            subject = columns[0].text
            marks = columns[1].text
            result[subject] = marks
    sgpa_web = driver.find_element("xpath","/html/body/div[3]/div[1]/div/div/div[3]/div[2]/div/div/div/div/div/div/div[7]")
    sgpa=sgpa_web.text
    
    reset = driver.find_element("xpath","/html/body/div[3]/div[1]/div/div/div[2]/div[2]/div[2]/div/div/div/div/form/div[4]/div/button[2]")
    reset.click()
    
    return (result,sgpa)

def convert_date(date):
    lst = date.split(".")
    return lst[0]+"/"+lst[1]+"/"+lst[2]


def get_screenshot():
    driver.execute_script("document.body.style.zoom='50%'")

    target_element = driver.find_element("xpath","/html/body/div[3]/div[1]")
    driver.execute_script("window.scrollBy(0, 500);")
    time.sleep(1)
    target_element.screenshot('temp.png')
    print("Cropping Image")
    crop = Image.open('temp.png')
    cropped_image = crop.crop((485, 78, 1040, 548))
    cropped_image.save('temp.png')
    crop.close()
    driver.execute_script("document.body.style.zoom='100%'")


def get_document(name,regno,dob):
    doc.add_heading(f'{regno}  {name}  ', 4)
    img_path = "temp.png"
    doc.add_picture(img_path, width=Inches(4.0))  
    
workbook = openpyxl.load_workbook('21-25 IT A.xlsx')
worksheet = workbook.active

idx=2
num_rows = worksheet.max_row

while num_rows>=idx:
    name=worksheet.cell(row=idx, column=1).value
    reg_no=worksheet.cell(row=idx,column=2).value
    dob = worksheet.cell(row=idx,column=3).value
    
    print(name,reg_no,dob) 
    # time.sleep(0.5)
    # get_result(reg_no,convert_date(dob))
    # time.sleep(0.5)
    # result,sgpa = read_result()
    # time.sleep(0.5)
    # get_screenshot()
    # time.sleep(0.5)
    # get_document(name,reg_no,dob) 
    # print(result,sgpa)

    try:
        time.sleep(0.5)
        get_result(reg_no,convert_date(dob))
        time.sleep(0.5)
        result,sgpa = read_result()
        time.sleep(0.5)
        get_screenshot()
        time.sleep(0.5)
        get_document(name,reg_no,dob) 
        print(result,sgpa)
    except Exception as e:
        idx += 1
        continue
    

    idx+=1
    print("\n\n")

print(idx)

doc.save('21-25 IT A Results.docx')
driver.close()